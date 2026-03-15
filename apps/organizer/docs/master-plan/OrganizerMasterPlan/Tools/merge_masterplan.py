# -*- coding: utf-8 -*-
"""
Merge and deduplicate Omni Organizer Master Plan .docx files.
Outputs:
- combined.docx: merged unique paragraphs
- combined.txt: merged unique paragraphs as plain text
- combine_report.json: details on duplicates and sources
- combined_outline.txt: extracted heading outline to understand structure
"""
import os
import re
import json
import hashlib
from collections import defaultdict

try:
    from docx import Document
except ImportError:
    raise SystemExit("python-docx not installed. Please run: python -m pip install python-docx")

DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = DIR  # same directory as script

OUTPUT_DOCX = os.path.join(INPUT_DIR, "combined.docx")
OUTPUT_TXT = os.path.join(INPUT_DIR, "combined.txt")
OUTPUT_REPORT = os.path.join(INPUT_DIR, "combine_report.json")
OUTPUT_OUTLINE = os.path.join(INPUT_DIR, "combined_outline.txt")

DOCX_PATTERN = re.compile(r"^omnia_master_plan_bolum.*\.docx$", re.IGNORECASE)

HEADING_STYLE_PREFIXES = ("Heading", "Baslik", "Başlık")  # common heading style names
HEADING_TEXT_PATTERN = re.compile(r"^(?:B[öo]l[üu]m|Chapter|Section)\s*\d+|^(?:Giri[sş]|Öns[öo]z|Conclusion|Sonu[cç])\b", re.IGNORECASE)

SECTION_NUM_PATTERNS = [
    re.compile(r"^(?:B[öo]l[üu]m)\s*(\d+)", re.IGNORECASE),
    re.compile(r"^(?:Chapter|Section)\s*(\d+)", re.IGNORECASE),
]

APPENDIX_TITLE_PATTERN = re.compile(r"^EKLER\b", re.IGNORECASE)
LONG_VISION_TITLE_PATTERN = re.compile(r"^UZUN\s+VADEL[İI]\s+V[İI]ZYON\b", re.IGNORECASE)


def extract_section_number(text: str):
    t = (text or '').strip()
    for pat in SECTION_NUM_PATTERNS:
        m = pat.search(t)
        if m:
            try:
                return int(m.group(1))
            except Exception:
                return None
    return None


def section_sort_key(key):
    kind, val = key
    order_map = {
        'front': -1,     # belge başlığı vb.
        'num': 0,        # Bölüm 1..N
        'tail': 1,       # numarasız genel içerik
        'vision': 2,     # UZUN VADELİ VİZYON
        'appendix': 3,   # EKLER
    }
    base = order_map.get(kind, 1)
    if kind == 'num':
        return (base, val)
    return (base, 0)


def list_docx_files(path):
    names = [n for n in os.listdir(path) if n.lower().endswith('.docx') and DOCX_PATTERN.match(n)]
    return [os.path.join(path, n) for n in names]


def normalize_text(text):
    # Collapse whitespace, strip, lower-case for dedup, but keep punctuation
    t = re.sub(r"\s+", " ", text or "").strip()
    return t.lower()


def paragraph_is_heading(p):
    text = (p.text or '').strip()
    if not text:
        return False, None
    style_name = getattr(getattr(p, 'style', None), 'name', '') or ''
    # Style-based heuristic
    for pref in HEADING_STYLE_PREFIXES:
        if style_name.startswith(pref):
            # Extract numeric level if present like Heading 1, Heading 2, else 1
            m = re.search(r"(\d+)$", style_name)
            level = int(m.group(1)) if m else 1
            return True, level
    # Text pattern heuristic
    if HEADING_TEXT_PATTERN.search(text):
        return True, None
    # All caps short lines might be headings
    if len(text) < 80 and text.isupper():
        return True, None
    return False, None


def read_paragraphs_with_meta(filepath):
    doc = Document(filepath)
    for p in doc.paragraphs:
        raw = (p.text or '')
        text = raw.strip()
        if not text:
            continue
        is_heading, level = paragraph_is_heading(p)
        style_name = getattr(getattr(p, 'style', None), 'name', '') or ''
        yield {
            'text': text,
            'normalized': normalize_text(text),
            'is_heading': is_heading,
            'heading_level': level,
            'style_name': style_name,
        }


def main():
    files = list_docx_files(INPUT_DIR)
    if not files:
        raise SystemExit("No .docx files found to process.")

    files_sorted = sorted(files)

    seen_hashes = set()
    kept = []  # list of dicts: {text, is_heading, heading_level, style_name, first_source, section_key}
    sources_map = defaultdict(list)  # norm_text_hash -> [filenames]

    def hash_norm(s):
        return hashlib.sha256(s.encode('utf-8')).hexdigest()

    for fp in files_sorted:
        base = os.path.basename(fp)
        current_section = ('front', 0)
        try:
            for item in read_paragraphs_with_meta(fp):
                if item['is_heading']:
                    # Bölüm numarası var mı?
                    secnum = extract_section_number(item['text'])
                    if secnum is not None:
                        current_section = ('num', secnum)
                    else:
                        # Özel başlıklar
                        t = item['text']
                        if APPENDIX_TITLE_PATTERN.search(t):
                            current_section = ('appendix', 0)
                        elif LONG_VISION_TITLE_PATTERN.search(t):
                            current_section = ('vision', 0)
                        elif current_section[0] == 'front':
                            # İlk numarasız ciddi başlıktan itibaren tail'e taşı
                            current_section = ('tail', 0)
                norm = item['normalized']
                if not norm:
                    continue
                h = hash_norm(norm)
                sources_map[h].append(base)
                if h in seen_hashes:
                    continue
                seen_hashes.add(h)
                kept.append({
                    'text': item['text'],
                    'is_heading': item['is_heading'],
                    'heading_level': item['heading_level'],
                    'style_name': item['style_name'],
                    'first_source': base,
                    'section_key': current_section,
                })
        except Exception as e:
            print(f"Warning: failed to read {base}: {e}")

    # Bölüm bazında grupla ve 1..N sırasına göre düzleştir
    buckets = defaultdict(list)
    for it in kept:
        sk = it.get('section_key', ('front', 0))
        buckets[sk].append(it)

    ordered_sections = sorted(buckets.keys(), key=section_sort_key)
    out_items = []
    for sk in ordered_sections:
        out_items.extend(buckets[sk])

    # Write combined.docx
    out_doc = Document()
    for item in out_items:
        if item['is_heading']:
            # Try to map heading levels to existing styles; fallback to normal
            level = item['heading_level'] or 1
            style = f"Heading {level}" if level and level in (1, 2, 3, 4) else None
            if style and style in [s.name for s in out_doc.styles]:
                out_doc.add_paragraph(item['text'], style=style)
            else:
                out_doc.add_paragraph(item['text'])
        else:
            out_doc.add_paragraph(item['text'])
    out_doc.save(OUTPUT_DOCX)

    # Write combined.txt
    with open(OUTPUT_TXT, 'w', encoding='utf-8') as f:
        for item in out_items:
            f.write(item['text'] + "\n\n")

    # Write outline
    with open(OUTPUT_OUTLINE, 'w', encoding='utf-8') as f:
        for item in out_items:
            if item['is_heading']:
                level = item['heading_level'] or 1
                indent = '  ' * (max(0, level - 1))
                f.write(f"{indent}- {item['text']}\n")

    # Write report
    duplicates = {h: srcs for h, srcs in sources_map.items() if len(srcs) > 1}
    report = {
        'input_files_order': [os.path.basename(p) for p in files_sorted],
        'unique_paragraphs': len(kept),
        'duplicate_paragraph_groups': len(duplicates),
        'duplicates_example': dict(list(duplicates.items())[:10]),
    }
    with open(OUTPUT_REPORT, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

    print("Done.")
    print(f"Files written:\n- {OUTPUT_DOCX}\n- {OUTPUT_TXT}\n- {OUTPUT_OUTLINE}\n- {OUTPUT_REPORT}")


if __name__ == '__main__':
    main()